import streamlit as st
import os
import requests
import tempfile
import zipfile
import re
import time
import io
import textwrap
import base64
import qrcode
# import pypandoc
import pythoncom
import urllib.parse
import streamlit.components.v1 as components
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE  # Importação corrigida
from docx2pdf import convert
from PIL import Image, ImageDraw, ImageFont
from datetime import datetime
# from pixqrcode import PixQrCode
# from pixqrcodegen import Payload


pythoncom.CoInitialize()


# Função para gerar BR Code (manual)
def gerar_br_code_pix(chave, nome, cidade, txid, valor=""):
    gui = "br.gov.bcb.pix"
    payload_format_indicator = "000201"
    merchant_account_info = f"26{len(gui + '01' + chave):02}00{len(gui):02}{gui}01{len(chave):02}{chave}"
    merchant_category_code = "52040000"
    transaction_currency = "5303986"
    transaction_amount = f"54{len(valor):02}{valor}" if valor else ""
    country_code = "5802BR"
    name = nome[:25]
    city = cidade[:15]
    merchant_name = f"59{len(name):02}{name}"
    merchant_city = f"60{len(city):02}{city}"
    txid_field = f"05{len(txid):02}{txid}"
    additional_data_field = f"62{len(txid_field):02}{txid_field}"
    crc_placeholder = "6304"

    full_data = (
        payload_format_indicator
        + merchant_account_info
        + merchant_category_code
        + transaction_currency
        + transaction_amount
        + country_code
        + merchant_name
        + merchant_city
        + additional_data_field
        + crc_placeholder
    )

    # Cálculo do CRC-16/CCITT-FALSE
    def crc16_ccitt(data: str) -> str:
        data = bytearray(data.encode("utf-8"))
        crc = 0xFFFF
        for byte in data:
            crc ^= byte << 8
            for _ in range(8):
                if crc & 0x8000:
                    crc = (crc << 1) ^ 0x1021
                else:
                    crc <<= 1
                crc &= 0xFFFF
        return format(crc, "04X")

    crc = crc16_ccitt(full_data)
    return full_data + crc


# Função para redimensionar logo
def redimensionar_logo(caminho_logo, largura_max_inch=2.5):
    img = Image.open(caminho_logo)
    largura, altura = img.size
    nova_largura_px = 300
    fator = nova_largura_px / largura
    nova_altura_px = int(altura * fator)
    img_redimensionada = img.resize((nova_largura_px, nova_altura_px))
    caminho_logo_temp = os.path.join(
        tempfile.gettempdir(), "logo_redimensionada.png")
    img_redimensionada.save(caminho_logo_temp)
    return caminho_logo_temp


# Função para ajustar bordas
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = 'w:{}'.format(edge)
            element = tcPr.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcPr.append(element)
            for key in ["val", "sz", "space", "color"]:
                element.set(qn('w:{}'.format(key)), str(edge_data.get(
                    key, 'single' if key == "val" else 8 if key == "sz" else 0 if key == "space" else "000000")))


# Função para gerar a primeira página
def gerar_primeira_pagina(doc, caminho_logo, empresa, area_negocio, lider_negocio, area_dev, lider_dev, nome_painel, data):

    # Obtém a data e hora atuais
    agora = datetime.now()
    # Extrai o ano
    ano_atual = agora.year

    # Função auxiliar que blinda a tabela inteira com bordas
    def blindar_tabela(table):
        for row in table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top={"val": "single"},
                    left={"val": "single"},
                    bottom={"val": "single"},
                    right={"val": "single"}
                )
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def fixar_layout_tabela(table):
        """
        Blinda o layout da tabela para que o Word não estique as colunas.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr

        tblLayout = OxmlElement('w:tblLayout')
        tblLayout.set(qn('w:type'), 'fixed')
        tblPr.append(tblLayout)

    if caminho_logo:
        try:
            # Redimensiona a logo
            image = Image.open(caminho_logo)
            max_width = 2 * 96  # 2 polegadas em pixels (96 dpi padrão)
            if image.width > max_width:
                w_percent = (max_width / float(image.width))
                h_size = int((float(image.height) * float(w_percent)))
                image = image.resize((max_width, h_size),
                                     Image.Resampling.LANCZOS)

        except Exception as e:
            st.warning(f"⚠️ Erro ao adicionar a logo: {e}")
            image = None

    else:
        image = None

    # Se a logo não foi carregada ou está vazia, cria uma imagem padrão
    if image is None:
        width_px = 192  # 2 polegadas * 96 dpi
        height_px = 96
        image = Image.new('RGB', (width_px, height_px), color=(230, 230, 230))
        draw = ImageDraw.Draw(image)

        # Fonte padrão
        font = ImageFont.load_default()

        text = "LOGO"
        bbox = draw.textbbox((0, 0), text, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]

        draw.text(
            ((width_px - text_width)/2, (height_px - text_height)/2),
            text,
            fill=(110, 110, 110),
            font=font
        )

    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)

    # Tabela Principal
    table = doc.add_table(rows=8, cols=7)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        row.height = Inches(0.1)

    # Aplica borda cegamente em toda a grade ANTES das mesclagens
    blindar_tabela(table)

    # --- Mesclagens e preenchimentos ---
    # Logo
    cell_logo_00 = table.cell(0, 0).merge(table.cell(0, 1))
    cell_logo_01 = table.cell(1, 0).merge(table.cell(1, 1))
    cell_logo_02 = table.cell(2, 0).merge(table.cell(2, 1))
    cell_logo_03 = table.cell(3, 0).merge(table.cell(3, 1))
    cell_logo_bloco = table.cell(0, 0).merge(table.cell(3, 0))

    paragraph = cell_logo_bloco.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(img_byte_arr, width=Inches(1.2))
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Relatório
    cell_relatorio = table.cell(0, 2).merge(table.cell(0, 3))
    cell_relatorio.text = "DOCUMENTAÇÃO DE RELATÓRIO"
    cell_relatorio.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for paragraph in cell_relatorio.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.bold = True

    # Número
    cell_numero = table.cell(0, 4).merge(table.cell(0, 6))
    cell_numero.width = Inches(2.5)

    for paragraph in cell_numero.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(1.25), WD_ALIGN_PARAGRAPH.CENTER)
        run_numero_label = paragraph.add_run("Nº:")
        run_numero_label.font.size = Pt(7)
        paragraph.add_run("\t")
        run_relatorio_numero = paragraph.add_run(
            f"RL-BI-{str(ano_atual)}-0001")
        run_relatorio_numero.font.size = Pt(10)
        run_relatorio_numero.font.bold = True

    # Segunda linha
    cell_empresa = table.cell(1, 2).merge(table.cell(1, 5))
    cell_empresa.width = Inches(3.5)
    cell_folha = table.cell(1, 6)
    cell_folha.width = Inches(1.5)

    for paragraph in cell_empresa.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(1.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_empresa_label = paragraph.add_run("EMPRESA:")
        run_empresa_label.font.size = Pt(7)
        paragraph.add_run("\t")
        run_empresaTxt_label = paragraph.add_run(str(empresa))
        run_empresaTxt_label.font.size = Pt(10)
        run_empresaTxt_label.font.bold = True

    for paragraph in cell_folha.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(0.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_folha_label = paragraph.add_run("FOLHAS:")
        run_folha_label.font.size = Pt(7)
        paragraph.add_run("\t")
        run_folhaTxt_label = paragraph.add_run("1 de ??")
        run_folhaTxt_label.font.size = Pt(10)
        run_folhaTxt_label.font.bold = True

    # Terceira linha
    cell_area_negocio = table.cell(2, 2).merge(table.cell(2, 5))
    cell_area_negocio.width = Inches(3.5)
    cell_lider_negocio = table.cell(2, 6)
    cell_lider_negocio.width = Inches(1.5)

    for paragraph in cell_area_negocio.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(1.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_area_negocio = paragraph.add_run("ÁREA DE NEGÓCIO:")
        run_area_negocio.font.size = Pt(7)
        paragraph.add_run("\t")
        run_area_negocioTxt_label = paragraph.add_run(str(area_negocio))
        run_area_negocioTxt_label.font.size = Pt(10)
        run_area_negocioTxt_label.font.bold = True

    for paragraph in cell_lider_negocio.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(0.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_lider_negocio_label = paragraph.add_run("LÍDER DE NEGÓCIO:")
        run_lider_negocio_label.font.size = Pt(7)
        paragraph.add_run("\n")
        run_lider_negocioTxt_label = paragraph.add_run(str(lider_negocio))
        run_lider_negocioTxt_label.font.size = Pt(10)
        run_lider_negocioTxt_label.font.bold = True

    # Quarta linha
    cell_area_dev = table.cell(3, 2).merge(table.cell(3, 5))
    cell_area_dev.width = Inches(3.5)
    cell_lider_dev = table.cell(3, 6)
    cell_lider_dev.width = Inches(1.5)

    for paragraph in cell_area_dev.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(1.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_area_dev_label = paragraph.add_run("ÁREA DE DESENVOLVIMENTO:")
        run_area_dev_label.font.size = Pt(7)
        paragraph.add_run("\t")
        run_area_devTxt_label = paragraph.add_run(str(area_dev))
        run_area_devTxt_label.font.size = Pt(10)
        run_area_devTxt_label.font.bold = True

    for paragraph in cell_lider_dev.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(0.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_lider_dev_label = paragraph.add_run("LÍDER DEV:")
        run_lider_dev_label.font.size = Pt(7)
        paragraph.add_run("\n")
        run_lider_devTxt_label = paragraph.add_run(str(lider_dev))
        run_lider_devTxt_label.font.size = Pt(10)
        run_lider_devTxt_label.font.bold = True

    # Quinta e Sexta linha
    table.cell(4, 0).merge(table.cell(5, 1))
    cell_titulo = table.cell(4, 2).merge(table.cell(5, 5))

    cell_divisao = table.cell(4, 6)
    cell_divisao.text = "TIC/PLAT/DSIG/SSIG"
    cell_divisao.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_tipo = table.cell(5, 6)
    cell_tipo.text = "INTERNA"
    cell_tipo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for paragraph in cell_titulo.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        paragraph_format = paragraph.paragraph_format
        paragraph_format.tab_stops.add_tab_stop(
            Inches(1.75), WD_ALIGN_PARAGRAPH.CENTER)
        run_titulo_label = paragraph.add_run("TÍTULO:")
        run_titulo_label.font.size = Pt(7)
        paragraph.add_run("\t")
        run_tituloTxt_label = paragraph.add_run(str(nome_painel))
        run_tituloTxt_label.font.size = Pt(10)
        run_tituloTxt_label.font.bold = True

    for paragraph in cell_divisao.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(9)
            run.font.bold = True

    for paragraph in cell_tipo.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(10)
            run.font.bold = True

    # Sétima linha
    table.cell(6, 0).merge(table.cell(6, 6))

    # Oitava linha
    cell_indice = table.cell(7, 0).merge(table.cell(7, 6))
    cell_indice.text = "ÍNDICE DE REVISÕES"
    cell_indice.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for paragraph in cell_indice.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(13)
            run.font.bold = True

    # Blinda novamente após mesclagens
    blindar_tabela(table)

    table_rev_desc = doc.add_table(rows=2, cols=7)
    table_rev_desc.autofit = False
    table_rev_desc.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Nona linha
    cell_rev = table_rev_desc.cell(0, 0)
    cell_rev.text = "REV."
    cell_rev.width = Inches(0.5)
    cell_rev.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_desc = table_rev_desc.cell(0, 1).merge(table_rev_desc.cell(0, 6))
    cell_desc.text = "DESCRIÇÃO E/OU FOLHAS ATINGIDAS"
    cell_desc.width = Inches(6)
    cell_desc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for paragraph in cell_desc.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(0)
        for run in paragraph.runs:
            run.font.size = Pt(10)
            # run.font.bold = True

    # Décima linha
    cell_valrev = table_rev_desc.cell(1, 0)
    cell_valrev.width = Inches(0.5)
    cell_valrev.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    cell_emissao = table_rev_desc.cell(1, 1).merge(table_rev_desc.cell(1, 6))
    cell_emissao.width = Inches(6)
    cell_emissao.text = "EMISSÃO ORIGINAL"

    # doc.add_paragraph()

    # Blinda novamente após mesclagens
    blindar_tabela(table_rev_desc)

    # Revisões
    table_rev = doc.add_table(rows=6, cols=7)
    table_rev.autofit = False
    table_rev.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Blinda o layout da tabela (ESSA LINHA É A CHAVE)
    fixar_layout_tabela(table_rev)

    col_width = Inches(0.93)

    # Força cada célula a ter largura fixa
    for row in table_rev.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = col_width

    # Define os headers corretos para 7 colunas
    headers = ["", "REV.", "REV. A", "REV. B", "REV. C", "REV. D", "REV. E"]
    for idx, rev in enumerate(headers):
        table_rev.cell(0, idx).text = rev
        table_rev.cell(0, idx).vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    labels = ["DATA", "EXECUÇÃO", "VERIFICAÇÃO", "APROVAÇÃO"]
    for i, label in enumerate(labels, start=1):
        cell = table_rev.cell(i, 0)
        cell.text = ""  # Limpa o conteúdo anterior
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(label)
        run.font.size = Pt(8)  # Define o tamanho da fonte

    # Adicionar a data na célula correspondente ("DATA" x "REV.")
    # +1 porque cabeçalho está na linha 0
    linha_data = labels.index("DATA") + 1
    coluna_rev = headers.index("REV.")

    cell_data = table_rev.cell(linha_data, coluna_rev)
    cell_data.text = ""  # Limpa conteúdo anterior
    paragraph_data = cell_data.paragraphs[0]
    run_data = paragraph_data.add_run(data.strftime("%d/%m/%Y"))
    run_data.font.size = Pt(8)

    # Rodapé
    cell_rodape = table_rev.cell(5, 0).merge(table_rev.cell(5, 6))
    cell_rodape.text = ""  # Limpa antes de formatar

    paragraph = cell_rodape.paragraphs[0]
    run = paragraph.add_run(
        f"DE ACORDO COM A NORMA INTERNA XYZ, AS INFORMAÇÕES DESTE DOCUMENTO SÃO PROPRIEDADE EXCLUSIVA DA {empresa.upper()} , SENDO PROIBIDA A UTILIZAÇÃO FORA DA SUA FINALIDADE. FORMULÁRIO PADRONIZADO PELA NORMA {empresa.upper()} X-123-ABC.Y.")
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(255, 0, 0)
    run.font.bold = True

    blindar_tabela(table_rev)

    doc.add_paragraph()

    # p = doc.add_paragraph()
    # run = p.add_run("DE ACORDO COM A D1-19BR-00337, AS INFORMAÇÕES DESTE DOCUMENTO SÃO PROPRIEDADE DA PETROBRAS, SENDO PROIBIDA A UTILIZAÇÃO FORA DA SUA FINALIDADE. FORMULÁRIO PADRONIZADO PELA NORMA PETROBRAS N-381-REV.M.")
    # font = run.font
    # font.color.rgb = RGBColor(255, 0, 0)
    # font.size = Pt(8)
    # p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    doc.add_page_break()


# Configuração de estilos personalizados
def configurar_estilos(doc):
    styles = doc.styles

    # Estilo para título principal (Paragraph Style)
    if 'TituloPrincipal' not in styles:
        titulo_style = styles.add_style(
            'TituloPrincipal', WD_STYLE_TYPE.PARAGRAPH)
        titulo_style.paragraph_format.space_after = Pt(12)

    # Estilo para subtítulos
    if 'Subtitulo' not in styles:
        subtitulo_style = styles.add_style(
            'Subtitulo', WD_STYLE_TYPE.PARAGRAPH)
        subtitulo_style.paragraph_format.space_before = Pt(18)
        subtitulo_style.paragraph_format.space_after = Pt(12)

    # Estilo para nível abaixo do subtítulo
    if 'Subsubtitulo' not in styles:
        subsubtitulo_style = styles.add_style(
            'Subsubtitulo', WD_STYLE_TYPE.PARAGRAPH)
        subsubtitulo_style.paragraph_format.space_before = Pt(18)
        subsubtitulo_style.paragraph_format.space_after = Pt(12)

    # Estilo para código M
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
        code_style.paragraph_format.left_indent = Inches(0.5)
        code_style.paragraph_format.space_before = Pt(6)
        code_style.paragraph_format.space_after = Pt(6)


def adicionar_titulo(doc, texto, estilo, espaco_before=28):
    paragrafo = doc.add_paragraph(style=estilo)
    run = paragrafo.add_run(texto)

    # Aplicação manual da fonte
    if estilo == 'TituloPrincipal':
        run.font.name = 'Petrobras Sans'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x0F, 0x0F, 0x0F)
        run.font.bold = True
        paragrafo.paragraph_format.space_after = Pt(14)

    elif estilo == 'Subtitulo':
        run.font.name = 'Petrobras Sans'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x0F, 0x0F)
        run.font.bold = True
        paragrafo.paragraph_format.space_before = Pt(espaco_before)
        paragrafo.paragraph_format.space_after = Pt(14)

    elif estilo == 'Subsubtitulo':
        run.font.name = 'Petrobras Sans'
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0F, 0x0F, 0x0F)
        run.font.bold = True
        paragrafo.paragraph_format.space_before = Pt(espaco_before)
        paragrafo.paragraph_format.space_after = Pt(14)

    return paragrafo


def encontrar_arquivos_tmdl(diretorio):
    tmdl_files = []
    for root, dirs, files in os.walk(diretorio):
        for file in files:
            if file.endswith(".tmdl"):
                tmdl_files.append(os.path.join(root, file))
    return tmdl_files


def extrair_metadados(conteudo):
    # Regex de tabelas
    table_pattern = re.compile(
        r"table\s+(?:'([^']+)'|(\w+))\s+lineageTag:\s*([\w-]+)",
        re.MULTILINE
    )

    # Regex de colunas
    column_pattern = re.compile(
        r"column\s+(?:'([^']+)'|(\w+))"  # Nome da coluna com ou sem aspas
        # Expressão com crase ou multiline
        r"(?:\s*=\s*(?:(```(?:.|\n)*?```)|((?:.|\n)*?)(?=\n\s+(?:dataType|formatString|lineageTag|summarizeBy|sortByColumn|annotation|column|variation|partition|$))))?"
        r"(?:\s*dataType:\s*(\w+))?"  # Tipo (opcional)
        r"(?:\s*summarizeBy:\s*(\w+))?"  # summarizeBy (opcional)
        # sortByColumn (opcional)
        r"(?:\s*sortByColumn:\s+(?:'([^']+)'|(\w+)))?",
        re.MULTILINE
    )

    # Regex de partition
    partition_pattern = re.compile(
        # <- captura 'calculated', 'query', etc.
        r"partition\s+(?:'([^']+)'|(\w+))\s*=\s*(\w+)\s*"
        r"mode:\s*(\w+)\s*"
        r"(?:queryGroup:\s*([^\n]+)\s*)?"
        r"(?:annotation\s+[^\n]+\s*)*"
        r"source\s*=\s*(?:(?:```([\s\S]+?)```)|(?:'''([\s\S]+?)''')|([\s\S]+?))"
        r"(?=\s*(?:partition|annotation|mode|queryGroup|source|\Z))",
        re.DOTALL
    )

    # Regex de medidas
    measure_pattern = re.compile(
        r"measure\s+(?:'([^']+)'|(\w+))\s*=\s*"
        r"(?:(```([\s\S]+?)```)|([\s\S]+?))"
        r"(?=\n(?:\s*(?:formatString|displayFolder|lineageTag|annotation|measure|table|partition|\Z)))",
        re.IGNORECASE
    )

    # Regex de expressões
    expression_pattern = re.compile(
        # Nome com ou sem aspas
        r"expression\s+(?:'([^']+)'|([^\s=]+))\s*=\s*"
        # Grupo 3: crase tripla
        r"(?:```([\s\S]+?)```"
        # Grupo 4: aspas triplas
        r"|\"\"\"([\s\S]+?)\"\"\""
        # Grupo 5: expressão normal
        r"|([\s\S]*?))"
        # Grupo 6: lineageTag
        r"\s*lineageTag:\s*([\w-]+)"
        # Grupo 7/8: queryGroup (opcional)
        r"(?:\s*queryGroup:\s*(?:'([^']+)'|([^\n]+)))?"
        # Ignorado (opcional)
        r"(?:\s*annotation\s+PBI_NavigationStepName\s*=\s*[^\n]+)?"
        r"\s*annotation\s+PBI_ResultType\s*=\s*([^\n]+)",    # Grupo 9: tipo
        re.IGNORECASE | re.DOTALL
    )

    # Regex de relacionamentos
    relationship_pattern = re.compile(
        r'relationship\s+([^\s]+)'                                 # (1) ID
        # (2) isActive (opcional)
        r'(?:\s+isActive:\s*(\w+))?'
        # (3) crossFilteringBehavior (opcional)
        r'(?:\s+crossFilteringBehavior:\s*(\w+))?'
        # ignora outros campos (ex: joinOnDateBehavior)
        r'(?:\s+\w+:\s*.*?)*?'
        # (4) fromTable, (5) fromColumn
        r'\s+fromColumn:\s*(?:["\']?)(.+?)["\']?\.(.+)'
        # (6) toTable, (7) toColumn
        r'\s+toColumn:\s*(?:["\']?)(.+?)["\']?\.(.+)',
        re.IGNORECASE
    )

    tabelas = []
    medidas = []
    expressoes = []
    relacionamentos = []

    for table_match in table_pattern.finditer(conteudo):
        nome_tabela = table_match.group(1) or table_match.group(2)
        if nome_tabela.startswith('LocalDateTable'):
            continue

        tabela = {
            'Nome': nome_tabela,
            'LineageTag': table_match.group(3),
            'Colunas': [],
            'Partition': {}
        }

        # Extraindo colunas da tabela
        for col_match in column_pattern.finditer(conteudo):
            expressao = (col_match.group(
                3) or col_match.group(4) or '').strip()
            if expressao.startswith('```') and expressao.endswith('```'):
                expressao = expressao[3:-3].strip()

            tabela['Colunas'].append({
                'Nome': col_match.group(1) or col_match.group(2),
                'Expressao': expressao,
                'Tipo': col_match.group(5) or '-',
                'summarizeBy': col_match.group(6) or '-',
                'sortByColumn': col_match.group(7) or col_match.group(8) or '-'
            })

        # Ordena alfabeticamente pelo nome das colunas (case-insensitive)
        tabela['Colunas'].sort(key=lambda m: m['Nome'].lower())

        # Extraindo partition
        for part_match in partition_pattern.finditer(conteudo):
            nome_particao = part_match.group(1) or part_match.group(2)
            if nome_particao == nome_tabela:
                tabela['Partition'] = {
                    'Tipo': part_match.group(3),
                    'Modo': part_match.group(4),
                    # agora seguro
                    'QueryGroup': (part_match.group(5) or '').strip(),
                    # agora seguro
                    'CodigoM': (part_match.group(6) or part_match.group(7) or part_match.group(8) or '').strip()
                }
                break

        tabelas.append(tabela)

    # Extraindo medidas
    for measure_match in measure_pattern.finditer(conteudo):
        nome_medida = measure_match.group(1) or measure_match.group(2)
        expressao = (measure_match.group(
            4) or measure_match.group(5) or '').strip()
        medidas.append({
            'Nome': nome_medida,
            'Expressao': expressao
        })

    # Ordena alfabeticamente pelo nome da medida (case-insensitive)
    medidas.sort(key=lambda m: m['Nome'].lower())

    # Extraindo expressões
    for expr_match in expression_pattern.finditer(conteudo):
        nome = expr_match.group(1) or expr_match.group(2)
        expressao = expr_match.group(3) or expr_match.group(
            4) or expr_match.group(5) or ''
        lineage_tag = expr_match.group(6)
        query_group = expr_match.group(7) or expr_match.group(8) or ''
        tipo = expr_match.group(9)

        expressoes.append({
            "Nome": nome.strip(),
            "Expressao": expressao.strip(),
            "LineageTag": lineage_tag.strip(),
            "QueryGroup": query_group.strip(),
            "Tipo": tipo.strip()
        })
    # Ordena alfabeticamente pelo nome da expressão (case-insensitive)
    expressoes.sort(key=lambda m: m['Nome'].lower())

    # Extraindo relacionamentos
    for rel_match in relationship_pattern.finditer(conteudo):
        id_rel = rel_match.group(1)
        is_active = rel_match.group(2)
        comportamento = rel_match.group(3)

        from_table = rel_match.group(4)
        from_col = rel_match.group(5)
        to_table = rel_match.group(6)
        to_col = rel_match.group(7)

        if "LocalDateTable" in from_table or "LocalDateTable" in to_table:
            continue

        relacionamentos.append({
            'ID': id_rel,
            'Ativo': "Não" if is_active == "false" else "Sim",
            'Comportamento': "Ambos" if comportamento == "bothDirections" else "Único",
            'Origem': f"{from_table}.{from_col}",
            'Destino': f"{to_table}.{to_col}"
        })

    # Ordena alfabeticamente pelo nome da tabela de origem (case-insensitive)
    relacionamentos.sort(key=lambda m: m['Origem'].lower())

    return tabelas, medidas, expressoes, relacionamentos


def criar_tabela_word(doc, titulo, cabecalhos, dados, estilo):
    configurar_estilos(doc)
    adicionar_titulo(doc, titulo, estilo)

    num_cols = len(cabecalhos)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False

    hdr_cells = table.rows[0].cells
    largura_total_max = 6.5  # polegadas úteis para página A4

    fator_caractere = 0.13  # 1 caractere ≈ 0.11 polegada

    larguras_minimas = []
    larguras_desejadas = []

    for i in range(num_cols):
        cab = cabecalhos[i]
        max_valor = max((len(str(linha[i]))
                        for linha in dados if i < len(linha)), default=0)

        largura_min = max(1.25, len(cab) * fator_caractere)
        largura_desejada = max(largura_min, max_valor * fator_caractere)

        larguras_minimas.append(largura_min)
        larguras_desejadas.append(largura_desejada)

    soma_desejada = sum(larguras_desejadas)

    if soma_desejada <= largura_total_max:
        larguras_finais = larguras_desejadas
    else:
        # ⚖️ Ajuste proporcional respeitando os mínimos dos cabeçalhos
        sobra_disponivel = largura_total_max - sum(larguras_minimas)
        pesos = [max(ld - lm, 0.01)
                 for ld, lm in zip(larguras_desejadas, larguras_minimas)]
        soma_pesos = sum(pesos)

        larguras_finais = [
            lm + (peso / soma_pesos) * sobra_disponivel
            for lm, peso in zip(larguras_minimas, pesos)
        ]

    # 🧩 Preencher cabeçalhos e aplicar largura
    for i, cabecalho in enumerate(cabecalhos):
        hdr_cells[i].text = cabecalho
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row in table.rows:
            row.cells[i].width = Inches(larguras_finais[i])

    # 🧩 Preencher dados
    for linha in dados:
        row_cells = table.add_row().cells
        for i, valor in enumerate(linha):
            row_cells[i].text = str(valor)
            row_cells[i].width = Inches(larguras_finais[i])

    return table


# Função para formatar código M
def formatar_codigo_m(doc, codigo):
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(codigo)

    # Fonte monoespaçada
    run.font.name = 'Courier New'  # ou 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(0, 0, 0)  # Preto puro

    # Estilo do parágrafo
    paragrafo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Adiciona fundo cinza claro e borda preta (bloco estilo código)
    p = paragrafo._p  # elemento XML do parágrafo
    pPr = p.get_or_add_pPr()

    # Fundo cinza claro
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), "EDEDED")  # cinza claro
    pPr.append(shd)

    # Borda preta em volta
    pBdr = OxmlElement('w:pBdr')

    for border_name in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')      # tamanho da borda
        border.set(qn('w:space'), '2')   # espaço da borda
        border.set(qn('w:color'), '000000')  # preto
        pBdr.append(border)

    pPr.append(pBdr)


# Função para formatar expressões DAX
def formatar_expressao_dax(nome, expressao):
    # Remove espaços extras
    expressao = expressao.strip()

    # Verifica se deve ficar em linha única
    if '\n' not in expressao and len(expressao.splitlines()) == 1:
        return f"{nome} = {expressao}"
    else:
        # Formatação multi-linha (só aplica quebras na expressão, não no nome)
        partes = []
        partes.append(f"{nome} = ")  # Nome sempre na primeira linha

        # Processa cada linha da expressão original
        for line in expressao.splitlines():
            line = line.strip()
            if line:  # Ignora linhas vazias
                # Aplica indentação e quebras apenas na expressão
                partes.append(f"    {line}")

        # Junta tudo
        expressao_formatada = "\n".join(partes)

        # Tratamento opcional para funções (descomente se necessário)
        # expressao_formatada = expressao_formatada.replace('(', '(\n        ')
        # expressao_formatada = expressao_formatada.replace(')', '\n    )')

        return expressao_formatada


# Função para gerar o documento Word
def adicionar_paragrafo_com_alinhamento(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    p.style.font.name = 'Arial'
    p.style.font.size = Pt(11)

    # Justifica apenas se for longo o suficiente
    if len(texto.split()) >= 10:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


# Função principal para gerar o documento
def gerar_documento(tabelas, medidas, expressoes, relacionamentos, lista_guias, opcao):
    doc = Document()
    configurar_estilos(doc)
    indice = 0

    texto_padrao = '''
    Lorem ipsum dolor sit amet, consectetur adipiscing elit, sed do eiusmod tempor incididunt ut labore et dolore magna aliqua. Ut enim ad minim veniam, quis nostrud exercitation ullamco laboris nisi ut aliquip ex ea commodo consequat.\n
    Duis aute irure dolor in reprehenderit in voluptate velit esse cillum dolore eu fugiat nulla pariatur. Excepteur sint occaecat cupidatat non proident, sunt in culpa qui officia deserunt mollit anim id est laborum.\n
    '''

    indice += 1

    # doc.add_heading('Documentação do Modelo Power BI', 0)

    if opcao == "Detalhada":

        adicionar_titulo(doc, f"{indice}.\tOBJETIVO", 'TituloPrincipal', 28)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run1 = p.add_run(texto_padrao)
        p.style.font.size = Pt(12)

        indice += 1
        doc.add_page_break()

        adicionar_titulo(doc, f"{indice}.\tVISÃO GERAL", 'TituloPrincipal', 28)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run1 = p.add_run(texto_padrao)
        p.style.font.size = Pt(12)

        indice += 1
        doc.add_page_break()

        adicionar_titulo(doc, f"{indice}.\tGUIAS", 'TituloPrincipal')
        for i, guia_nome in enumerate(lista_guias, start=1):
            # Adiciona título para cada guia
            adicionar_titulo(
                doc, f"{indice}.{i}.\tGUIA ({guia_nome})", 'Subtitulo')

            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run1 = p.add_run(texto_padrao)
            p.style.font.size = Pt(12)
            p.style.font.bold = False
            p.paragraph_format.line_spacing = Pt(11)

        indice += 1
        doc.add_page_break()

    if tabelas:

        adicionar_titulo(doc, f"{indice}.\tTABELAS", 'TituloPrincipal')

        for i, tabela in enumerate(tabelas, start=1):
            # Adiciona título para cada tabela
            # Se for a primeira tabela, usa o estilo 'Subtitulo' normal
            # Se não, usa o estilo 'Subtitulo' com tamanho 28
            if i == 1:
                adicionar_titulo(
                    doc, f"{indice}.{i}.\tTABELA ({tabela['Nome']})", 'Subtitulo')
            else:
                adicionar_titulo(
                    doc, f"{indice}.{i}.\tTABELA ({tabela['Nome']})", 'Subtitulo', 28)

            # p = doc.add_paragraph()
            # run1 = p.add_run("LineageTag: ")
            # run1.bold = True
            # run2 = p.add_run(tabela['LineageTag'])
            # run2.italic = True
            # p.paragraph_format.line_spacing = Pt(11)

            if tabela['Partition']:
                # Modo
                p = doc.add_paragraph()
                run1 = p.add_run("Modo: ")
                run1.bold = True
                run2 = p.add_run(tabela['Partition']['Modo'] or '-')
                run2.italic = True
                p.paragraph_format.line_spacing = Pt(11)

                # Tipo
                p = doc.add_paragraph()
                run1 = p.add_run("Tipo: ")
                run1.bold = True
                run2 = p.add_run(tabela['Partition']['Tipo'] or '-')
                run2.italic = True
                p.paragraph_format.line_spacing = Pt(11)

                # Query Group
                p = doc.add_paragraph()
                run1 = p.add_run("Query Group: ")
                run1.bold = True
                run2 = p.add_run(tabela['Partition']['QueryGroup'] or '-')
                run2.italic = True
                p.paragraph_format.line_spacing = Pt(11)

                # Código M / Expressão DAX
                p = doc.add_paragraph()
                texto_run1 = 'Código M:' if tabela['Partition']['Tipo'] == 'm' else 'Expressão DAX:'
                run1 = p.add_run(texto_run1)
                run1.bold = True
                p.paragraph_format.line_spacing = Pt(11)
                if texto_run1 == 'Código M:':
                    # Não adiciona o nome da tabela se for código M
                    formatar_codigo_m(doc, tabela['Partition']['CodigoM'])
                else:
                    # Adiciona o nome da tabela se for expressão DAX
                    codigo_m = tabela['Partition']['CodigoM']
                    prefixo = f"{tabela['Nome']} ="
                    separador = "\n" if '\n' in codigo_m else " "
                    formatar_codigo_m(doc, f"{prefixo}{separador}{codigo_m}")

            # Colunas
            if tabela['Colunas']:
                dados_colunas = [
                    [col['Nome'], col['Tipo'],
                        f"{col['Nome']} = {col['Expressao']}" if col['Expressao'] else '-']
                    for col in tabela['Colunas']
                ]
                criar_tabela_word(
                    doc,
                    f"{indice}.{i}.1\tColunas",
                    ['Nome', 'Tipo', 'Expressão'],
                    dados_colunas,
                    'Subsubtitulo'
                )

        doc.add_page_break()

    if medidas:
        # Preparar dados com formatação DAX
        indice += 1
        dados_medidas = []
        for med in medidas:
            expressao_formatada = formatar_expressao_dax(
                med['Nome'], med['Expressao'])
            dados_medidas.append([
                med['Nome'],
                expressao_formatada
            ])

        criar_tabela_word(
            doc,
            f"{indice}.\tMEDIDAS DAX",
            ['Medida', 'Expressão'],
            dados_medidas,
            'TituloPrincipal'
        )

        doc.add_page_break()

    if expressoes:
        # doc.add_paragraph('Funções e Parâmetros', style='Subtitulo')
        indice += 1
        adicionar_titulo(
            doc, f"{indice}.\tFUNÇÕES, PARÂMETROS E TABELAS (SEM CARGA)", 'TituloPrincipal')

        for expr in expressoes:
            for label, valor in [
                ('Nome', expr['Nome']),
                ('Tipo', expr['Tipo']),
                ('Query Group', expr['QueryGroup'])
            ]:
                par = doc.add_paragraph()
                run_label = par.add_run(f"{label}: ")
                run_label.bold = True
                run_valor = par.add_run(valor if valor else '-')
                run_valor.italic = True
                par.paragraph_format.line_spacing = Pt(8)

            # Expressão (em bloco separado)
            par_expr = doc.add_paragraph()
            texto_run_label = 'Código M:' if expr['Tipo'] == 'Table' else 'Expressão:'
            run_label = par_expr.add_run(texto_run_label)
            run_label.bold = True
            par_expr.paragraph_format.line_spacing = Pt(8)
            formatar_codigo_m(doc, expr['Expressao'] or '')

            doc.add_paragraph('')  # Espaço extra entre registros

        doc.add_page_break()

    if relacionamentos:
        indice += 1
        dados_rel = [
            [rel['Origem'], rel['Destino'],
                rel['Comportamento'], rel['Ativo']]
            for rel in relacionamentos
        ]
        criar_tabela_word(
            doc,
            f"{indice}.\tRELACIONAMENTOS",
            ['Origem', 'Destino', 'Comportamento', 'Ativo'],
            dados_rel,
            'TituloPrincipal'
        )

    return doc


# Função para processar o projeto e gerar o documento
def processar_projeto(diretorio, lista_guias, opcao):
    tmdl_files = encontrar_arquivos_tmdl(diretorio)
    if not tmdl_files:
        return None

    todas_tabelas = []
    todas_medidas = []
    todas_expressoes = []
    todos_relacionamentos = []

    for file_path in tmdl_files:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                conteudo = f.read()
                tabelas, medidas, expressoes, relacionamentos = extrair_metadados(
                    conteudo)
                todas_tabelas.extend(tabelas)
                todas_medidas.extend(medidas)
                todas_expressoes.extend(expressoes)
                todos_relacionamentos.extend(relacionamentos)
        except Exception as e:
            st.warning(
                f"Erro ao processar {os.path.basename(file_path)}: {str(e)}")
            continue

    if not (todas_tabelas or todas_medidas or todas_expressoes or todos_relacionamentos):
        return None

    return gerar_documento(todas_tabelas, todas_medidas, todas_expressoes, todos_relacionamentos, lista_guias, opcao)


# Função para exibir uma linha separadora estilizada
def linha_separadora(estilo="default"):
    if estilo == "sombra":
        html = '<hr style="border: none; height: 1px; background-color: #e0e0e0; box-shadow: 0 1px 2px rgba(0,0,0,0.1); margin: 25px 0;">'
    elif estilo == "colorida":
        html = '<hr style="border: none; height: 4px; background: linear-gradient(to right, #e01d1d, #dba20f); border-radius: 3px; margin: 25px 0;">'
    elif estilo == "tracejada":
        html = '<hr style="border: none; border-top: 1px dashed #bbb; margin: 25px 0;">'
    else:
        html = '<hr style="border: none; height: 2px; background-color: #ccc; margin: 40px 0;">'

    st.markdown(html, unsafe_allow_html=True)


# Função para atualizar o número de guias
def atualizar_guis():
    st.session_state.num_guis = st.session_state.num_input
    st.session_state.lista_guias = st.session_state.lista_guias[:st.session_state.num_guis]


# __  .__   __. .___________. _______ .______       _______    ___       ______  _______         _______.___________..______       _______     ___      .___  ___.  __       __  .___________.
# |  | |  \ |  | |           ||   ____||   _  \     |   ____|  /   \     /      ||   ____|       /       |           ||   _  \     |   ____|   /   \     |   \/   | |  |     |  | |           |
# |  | |   \|  | `---|  |----`|  |__   |  |_)  |    |  |__    /  ^  \   |  ,----'|  |__         |   (----`---|  |----`|  |_)  |    |  |__     /  ^  \    |  \  /  | |  |     |  | `---|  |----`
# |  | |  . `  |     |  |     |   __|  |      /     |   __|  /  /_\  \  |  |     |   __|         \   \       |  |     |      /     |   __|   /  /_\  \   |  |\/|  | |  |     |  |     |  |
# |  | |  |\   |     |  |     |  |____ |  |\  \----.|  |    /  _____  \ |  `----.|  |____    .----)   |      |  |     |  |\  \----.|  |____ /  _____  \  |  |  |  | |  `----.|  |     |  |
# |__| |__| \__|     |__|     |_______|| _| `._____||__|   /__/     \__\ \______||_______|   |_______/       |__|     | _| `._____||_______/__/     \__\ |__|  |__| |_______||__|     |__|
# --- Configuração da página ---
st.set_page_config(page_title="Documentador Power BI",
                   page_icon="🧪", layout="wide")
st.title("🧪📄 LabiDocs - Documentador Power BI")

# --- Estados iniciais ---
for key in [
    "logo_uploaded", "projeto_sem_logo", "arquivo_uploaded",
    "arquivo_valido", "gerando_documentacao", "documentacao_gerada",
    "logo_file_anterior", "projeto_sem_logo_anterior",
    "arquivo_zip_anterior", "docx_bytes", "pdf_bytes", "zip_path"
]:
    if key not in st.session_state:
        st.session_state[key] = None if "bytes" in key or "zip_path" in key else False

# --- Sidebar ---
sidebar_logo_placeholder = st.sidebar.empty()
st.sidebar.subheader("🔍 Modo de Documentação")
opcao = st.sidebar.selectbox("Selecione uma opção:", [
                             "Complementar", "Detalhada"])

# --- Estilo do Upload ---
st.markdown("""
<style>
    .stFileUploader > div > div {
        border: 2px dashed #4CAF50;
        border-radius: 8px;
        padding: 30px;
        text-align: center;
        background-color: #f8f9fa;
    }
    .stFileUploader > div > div:hover {
        border-color: #2B546B;
        background-color: #f0f3f5;
    }
</style>
""", unsafe_allow_html=True)

# --- Uploads ---
col1, col2 = st.columns(2)

# Upload da Logo
with col1:
    logo_file = st.file_uploader(
        "🖼️ Envie a logo do projeto", type=["png", "jpg", "jpeg"])
    if logo_file:
        st.session_state.logo_uploaded = True
        st.session_state.logo_file_bytes = logo_file.getvalue()
        sidebar_logo_placeholder.image(
            logo_file, use_container_width=True, output_format="PNG")
        st.success("✅ Logo carregada com sucesso!")
        st.session_state.projeto_sem_logo = False
    else:
        st.session_state.logo_uploaded = False
        st.checkbox("Projeto sem logo", key="projeto_sem_logo")
        if st.session_state.projeto_sem_logo:
            st.success("✔️ Continuando sem logo.")
        else:
            st.info("⬆️ Carregue uma logo para iniciar.")

# Reset de estados se a logo/projeto mudou
if (logo_file != st.session_state.logo_file_anterior or
        st.session_state.projeto_sem_logo != st.session_state.projeto_sem_logo_anterior):
    st.session_state.logo_file_anterior = logo_file
    st.session_state.projeto_sem_logo_anterior = st.session_state.projeto_sem_logo
    st.session_state.documentacao_gerada = False
    st.session_state.docx_bytes = None
    st.session_state.pdf_bytes = None

# Upload do arquivo .zip
with col2:
    uploaded_file = st.file_uploader(
        "📁 Carregue o arquivo .zip extraído do .pbix",
        type=["zip"],
        help="Selecione o arquivo .zip exportado do Power BI Desktop (PBIP)"
    )
    if uploaded_file:
        if uploaded_file != st.session_state.arquivo_zip_anterior:
            st.session_state.arquivo_zip_anterior = uploaded_file
            st.session_state.arquivo_uploaded = True
            st.session_state.documentacao_gerada = False
            st.session_state.docx_bytes = None
            st.session_state.pdf_bytes = None

            # Salva ZIP temporariamente
            tmpdir_upload = tempfile.gettempdir()
            zip_path_upload = os.path.join(tmpdir_upload, "projeto.zip")
            with open(zip_path_upload, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.session_state.zip_path = zip_path_upload

            # Verifica se contém .tmdl
            with tempfile.TemporaryDirectory() as tmpdir:
                with zipfile.ZipFile(st.session_state.zip_path, "r") as zip_ref:
                    zip_ref.extractall(tmpdir)
                tmdl_encontrado = any(
                    file.endswith(".tmdl") for root, _, files in os.walk(tmpdir) for file in files
                )
                st.session_state.arquivo_valido = tmdl_encontrado
                if tmdl_encontrado:
                    st.success("✅ Arquivo extraído com sucesso!")
                else:
                    st.warning(
                        "⚠️ O arquivo ZIP não contém .tmdl. Verifique a exportação completa do PBIP.")
    else:
        st.session_state.arquivo_uploaded = False
        st.info("⬆️ Carregue um arquivo .zip para iniciar.")

# --- Input de guias ---
if "num_guis" not in st.session_state:
    st.session_state.num_guis = 1
if "lista_guias" not in st.session_state:
    st.session_state.lista_guias = [""] * st.session_state.num_guis


if opcao == "Detalhada":
    st.markdown("---")
    st.subheader("📝 Detalhes do Painel")
    col1, col2, col3 = st.columns([2, 2, 2])

    with col1:
        nome_painel = st.text_input("📊 Nome do Painel")
        area_negocio = st.text_input("🌍 Área de Negócio")
        lider_negocio = st.text_input("👔 Líder - Área de Negócio")

    with col2:
        empresa = st.text_input("🏢 Empresa")  # Alterado o ícone de Empresa
        area_dev = st.text_input("💻 Área de Desenvolvimento")
        lider_dev = st.text_input("🧑‍💼 Líder - Área de Desenvolvimento")

    with col3:
        data = st.date_input("📅 Data")
        st.number_input(
            "Quantas guias tem o seu painel?",
            min_value=0, max_value=20,
            value=st.session_state.num_guis,
            key="num_input",
            step=1,
            on_change=atualizar_guis
        )
        for i in range(st.session_state.num_guis):
            valor = st.text_input(f"Nome da Guia {i + 1}", key=f"guia_{i}")
            if len(st.session_state.lista_guias) < st.session_state.num_guis:
                st.session_state.lista_guias.append(valor)
            else:
                st.session_state.lista_guias[i] = valor

    campos_preenchidos = all([
        nome_painel.strip(), area_negocio.strip(), lider_negocio.strip(),
        empresa.strip(), area_dev.strip(), lider_dev.strip(),
        bool(data),
        st.session_state.num_guis > 0,
        all(nome.strip() for nome in st.session_state.lista_guias)
    ])

    if not campos_preenchidos:
        st.warning(
            "⚠️ **Todos os campos são de preenchimento obrigatório.** Preencha-os para continuar.")
    else:
        st.success("✅ Todos os campos obrigatórios foram preenchidos!")
else:
    campos_preenchidos = True
    nome_painel = "Painel Power BI"

# --- Disponibilidade do botão ---
disponivel = (
    st.session_state.arquivo_valido
    and (st.session_state.logo_uploaded or st.session_state.projeto_sem_logo)
    and st.session_state.arquivo_uploaded
    and campos_preenchidos
)

if disponivel:
    st.success("📄 Tudo pronto para gerar a documentação!")
    botao_texto = "🚀 Gerar Documentação" if not st.session_state.gerando_documentacao else "⏳ Processando..."
    gerar_btn = st.sidebar.button(
        botao_texto, disabled=st.session_state.gerando_documentacao)

    if gerar_btn:
        st.session_state.gerando_documentacao = True
        st.session_state.documentacao_gerada = False
        st.session_state.docx_bytes = None
        st.session_state.pdf_bytes = None

        with tempfile.TemporaryDirectory() as tmpdir:
            extract_dir = os.path.join(tmpdir, "extraido")
            with zipfile.ZipFile(st.session_state.zip_path, 'r') as zip_ref:
                zip_ref.extractall(extract_dir)

            logo_path = io.BytesIO(
                st.session_state.logo_file_bytes) if st.session_state.logo_uploaded else None

            with st.sidebar.status('⏳ Processando arquivos e gerando documentação...', expanded=True):
                progresso = st.sidebar.progress(0, text="Iniciando...")
                time.sleep(0.5)

                doc = Document()
                progresso.progress(10, text="Carregando arquivos...")

                if opcao == "Detalhada":
                    gerar_primeira_pagina(doc, logo_path, empresa, area_negocio,
                                          lider_negocio, area_dev, lider_dev, nome_painel, data)

                doc_processado = processar_projeto(
                    extract_dir, st.session_state.lista_guias, opcao)

                if doc_processado:
                    for element in doc_processado.element.body:
                        doc.element.body.append(element)

                progresso.progress(30, text="Analisando dados...")
                time.sleep(1)

                if opcao != "Detalhada" and doc_processado:
                    for element in doc_processado.element.body:
                        doc.element.body.append(element)

                progresso.progress(60, text="Gerando documentação...")
                output_docx = os.path.join(
                    tmpdir, f"Documentação - {nome_painel}.docx")
                doc.save(output_docx)
                output_pdf = os.path.join(
                    tmpdir, f"Documentação - {nome_painel}.pdf")
                # convert(output_docx, output_pdf)

                try:
                    # A mágica acontece aqui!
                    # pypandoc.convert_file(
                    #   output_docx, 'pdf', outputfile=output_pdf)

                    with open(output_docx, "rb") as f_docx, open(output_pdf, "rb") as f_pdf:
                        st.session_state.docx_bytes = f_docx.read()
                        st.session_state.pdf_bytes = f_pdf.read()

                except Exception as e:
                    # st.error(f"Ocorreu um erro ao converter para PDF: {e}")
                    # st.info("O download do arquivo DOCX ainda está disponível.")
                    # Mesmo com erro no PDF, ainda disponibiliza o DOCX
                    with open(output_docx, "rb") as f_docx:
                        st.session_state.docx_bytes = f_docx.read()
                    st.session_state.pdf_bytes = None  # Garante que o botão de PDF não apareça

                progresso.progress(100, text="Finalizando...")
                time.sleep(0.5)

        st.session_state.documentacao_gerada = True
        st.session_state.gerando_documentacao = False

# --- Botões de download ---
if st.session_state.documentacao_gerada:
    st.sidebar.success("✅ Documentação gerada com sucesso!")
    st.sidebar.download_button(
        "📥 Baixar DOCX",
        st.session_state.docx_bytes,
        f"Documentação - {nome_painel}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    if st.session_state.pdf_bytes is not None:
        # INDENTE O BOTÃO PARA QUE FIQUE DENTRO DO NOVO "IF"
        st.sidebar.download_button(
            "📥 Baixar PDF",
            st.session_state.pdf_bytes,
            f"Documentação - {nome_painel}.pdf",
            mime="application/pdf"
        )


# Rodapé com PIX
if st.session_state.documentacao_gerada:
    chave_pix = "(21)96848-5316"
    codigo_pix_payload = "00020126360014BR.GOV.BCB.PIX0114+55219684853165204000053039865802BR5914LEANDRO TELLES6014RIO DE JANEIRO62070503***63041C67"
    img = qrcode.make(codigo_pix_payload)
    buffer = io.BytesIO()
    img.save(buffer, format="PNG")
    qr_base64 = base64.b64encode(buffer.getvalue()).decode()

    components.html(f"""
    <footer style="position: fixed; bottom: 0; left: 0; width: 100%; padding: 16px 24px;
        background: linear-gradient(90deg, #fdfbfb 0%, #ebedee 100%); border-top: 1px solid #ddd;
        display: flex; align-items: center; justify-content: center; flex-wrap: wrap; font-family: 'Segoe UI';
        box-shadow: 0 -1px 4px rgba(0,0,0,0.06); z-index: 9999;">

        <div style="margin-right: 20px; text-align: center;">
            <img src="data:image/png;base64,{qr_base64}" width="100" height="100" alt="QR PIX" style="border: 3px solid #ccc; border-radius: 10px;"/>
            <div style="font-size: 12px; color: #666; margin-top: 4px;">Escaneie com o app do banco</div>
        </div>

        <div style="max-width: 400px; color: #333; font-size: 14px; line-height: 1.5;">
            <p style="margin: 4px 0 0; font-size: 16px;"><strong>💡 Gostou da ferramenta?</strong></p>
            <p style="margin: 4px 0;">Se te ajudou, considere apoiar com uma contribuição via PIX:</p>
            <p id="pix-key" style="margin: 6px 0; background: #f2f2f2; padding: 6px 10px; border-radius: 6px;
                font-family: monospace; color: #222;">
                📬 <strong>Chave:</strong> {chave_pix}
            </p>
            <button onclick="copyPix()" style="margin-top: 4px; background: #4CAF50; color: white; border: none;
                padding: 6px 12px; border-radius: 6px; cursor: pointer; font-size: 13px;">📋 Copiar chave</button>
        </div>

        <div id="toast" style="visibility: hidden; min-width: 180px; background-color: #4CAF50; color: white;
            text-align: center; border-radius: 6px; padding: 10px; position: fixed; z-index: 10000;
            bottom: 100px; right: 30px; font-size: 14px; box-shadow: 0 2px 8px rgba(0,0,0,0.2);
            transition: visibility 0s, opacity 0.5s ease-in-out; opacity: 0;">✅ Chave PIX copiada!</div>

        <script>
            function copyPix() {{
                const text = document.getElementById("pix-key").innerText.replace("📬 Chave:", "").trim();
                navigator.clipboard.writeText(text).then(function() {{
                    var toast = document.getElementById("toast");
                    toast.style.visibility = "visible";
                    toast.style.opacity = "1";
                    setTimeout(function() {{
                        toast.style.opacity = "0";
                        toast.style.visibility = "hidden";
                    }}, 2500);
                }});
            }}
        </script>
    </footer>
    """, height=220)


pythoncom.CoUninitialize()
